from __future__ import annotations

import hashlib
import io
import re
from pathlib import Path

import fitz
from docx import Document as DocxDocument
from fastapi import HTTPException, UploadFile

from .config import settings


ALLOWED_EXTENSIONS = {".pdf", ".docx", ".txt"}


def _normalise_text(text: str) -> str:
    text = text.replace("\x00", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _extract_pdf(data: bytes) -> str:
    pages: list[str] = []
    try:
        with fitz.open(stream=data, filetype="pdf") as pdf:
            for index, page in enumerate(pdf, start=1):
                page_text = page.get_text("text").strip()
                if page_text:
                    pages.append(f"[Page {index}]\n{page_text}")
    except Exception as exc:
        raise HTTPException(status_code=422, detail=f"The PDF could not be read: {exc}") from exc
    return "\n\n".join(pages)


def _extract_docx(data: bytes) -> str:
    try:
        doc = DocxDocument(io.BytesIO(data))
    except Exception as exc:
        raise HTTPException(status_code=422, detail=f"The DOCX file could not be read: {exc}") from exc

    chunks: list[str] = []
    for index, paragraph in enumerate(doc.paragraphs, start=1):
        value = paragraph.text.strip()
        if not value:
            continue
        style = paragraph.style.name if paragraph.style else "Paragraph"
        chunks.append(f"[{style}, paragraph {index}]\n{value}")

    for table_index, table in enumerate(doc.tables, start=1):
        rows: list[str] = []
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(cells):
                rows.append(" | ".join(cells))
        if rows:
            chunks.append(f"[Table {table_index}]\n" + "\n".join(rows))

    return "\n\n".join(chunks)


def _extract_txt(data: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise HTTPException(status_code=422, detail="The text file encoding is not supported.")


async def read_and_extract(upload: UploadFile) -> tuple[bytes, str, str, str]:
    filename = Path(upload.filename or "upload").name
    extension = Path(filename).suffix.lower()
    if extension not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=415,
            detail="Only PDF, DOCX and TXT files are supported in this MVP.",
        )

    max_bytes = settings.max_upload_mb * 1024 * 1024
    data = await upload.read(max_bytes + 1)
    if len(data) > max_bytes:
        raise HTTPException(
            status_code=413,
            detail=f"The file exceeds the {settings.max_upload_mb} MB upload limit.",
        )
    if not data:
        raise HTTPException(status_code=422, detail="The uploaded file is empty.")

    if extension == ".pdf":
        text = _extract_pdf(data)
    elif extension == ".docx":
        text = _extract_docx(data)
    else:
        text = _extract_txt(data)

    text = _normalise_text(text)
    word_count = len(re.findall(r"\b\w+\b", text))
    if word_count < 250:
        raise HTTPException(
            status_code=422,
            detail=(
                "The document contains too little readable text to produce a reliable "
                "20-question ownership assessment. Scanned PDFs need OCR before upload."
            ),
        )

    digest = hashlib.sha256(data).hexdigest()
    return data, text, filename, digest


def build_context(text: str, limit: int | None = None) -> str:
    """Keep broad coverage when a document is longer than the model context budget."""
    limit = limit or settings.max_context_chars
    if len(text) <= limit:
        return text

    blocks = [block.strip() for block in text.split("\n\n") if block.strip()]
    if not blocks:
        return text[:limit]

    target_segments = 12
    segment_budget = max(1500, limit // target_segments)
    selected: list[str] = []
    for segment in range(target_segments):
        position = round(segment * (len(blocks) - 1) / max(1, target_segments - 1))
        start = max(0, position - 1)
        piece = "\n\n".join(blocks[start : start + 3])
        selected.append(piece[:segment_budget])

    context = "\n\n".join(selected)
    return context[:limit]
